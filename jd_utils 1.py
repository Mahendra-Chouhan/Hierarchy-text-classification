from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from pdfrw import PdfReader, PdfWriter
import cv2
from spire.doc import *
from spire.doc.common import *
import comtypes.client
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor  
from docx.shared import Inches, Pt, Cm
from datetime import datetime
# used to put date into "Month YYYY to Month YYYY" format and returns string.
def date_range(start_date, end_date):
    

    # Extract month and year from start and end dates
    start_month_year = start_date.strftime("%B %Y")
    end_month_year = end_date.strftime("%B %Y")

    # Return date range string
    return f"{start_month_year} to {end_month_year}"

# write given list of text into image and returns image path
def write_text_center(image_pth, text): 
    font=cv2.FONT_HERSHEY_SIMPLEX
    font_scale=1
    color=(0, 0, 0)
    thickness=2
    image = cv2.imread(image_pth)

    #resize image
    W= int(2480/2.92)
    H= int(3508/3.2)
    image=cv2.resize(image,(W,H))
    # Get the size of the text
    text_size = cv2.getTextSize(text[0], font, font_scale, thickness)[0]
    # Calculate the position to place the text at the center
    text_x = (image.shape[1] - text_size[0]) // 2
    text_y = (image.shape[0] + text_size[1]) // 2
    # Write the text on the image
    cv2.putText(image, text[0], (text_x, text_y), font, font_scale, color, thickness, lineType=cv2.LINE_AA)
    
    text_size_1 = cv2.getTextSize(text[1], font, font_scale, thickness)[0]
    # Calculate the position to place the text at the center
    text_x_1 = (image.shape[1] - text_size_1[0]) // 2
    text_y_1 = (image.shape[0] + text_size_1[1]) // 2

    cv2.putText(image, text[1], (text_x_1, text_y_1+50), font, font_scale, color, thickness, lineType=cv2.LINE_AA)
    filename=r"./front_img.png"
    cv2.imwrite(filename, image)
    return filename

#Removes first page of pdf and returns modified pdf's path
def remove_page_(pdf_path):
    # Define the reader and writer objects
    reader_input = PdfReader(pdf_path)
    writer_output = PdfWriter()

    # Go through the pages one after the next
    for current_page in range(len(reader_input.pages)):
        if current_page != 0:# Removing "0" th page
            writer_output.addpage(reader_input.pages[current_page])
    # Write the modified content to disk
    writer_output.write(pdf_path)

    return pdf_path

#Add's background colour to pdf and returns modified pdf's path
def add_offset_colour(doc_path):
    
    # Create a Document object
    document = Document()
    # Load a Word document
    document.LoadFromFile(doc_path)
    # Get the document's background
    background = document.Background
    # Set the background type as Color
    background.Type = BackgroundType.Color
    # Set the background color
    background.Color = Color.get_AliceBlue()

    #save the resulting document
    document.SaveToFile(doc_path, FileFormat.Docx2016)
    return doc_path

#convers document to PDF and return pdf's path
def convert_word_to_pdf(word_path, pdf_path):
   comtypes.CoInitialize()  # Initialize the COM library
   try:
       # Save the Word document as a PDF using Microsoft Word
       word = comtypes.client.CreateObject("Word.Application")
       docx_path = os.path.abspath(word_path)
       pdf_path = os.path.abspath(pdf_path)
       pdf_format = 17  # PDF file format code
       word.Visible = False
       in_file = word.Documents.Open(docx_path)
       in_file.SaveAs(pdf_path, FileFormat=pdf_format)
       in_file.Close()
   finally:
       word.Quit()
       comtypes.CoUninitialize()  # Uninitialize the COM library 

#Makes reports/documents first page and sets some parameters of remaining pages returns "docx" object
def make_document_first_page(doc,image_path,start_date,end_date,Message):

    range_date=date_range(start_date,end_date)
    myMessage = [Message,range_date]
    image_path=write_text_center(image_path,myMessage)

    sec=doc.add_section()# this page is required and intentionally left blank.
    new_section = doc.sections[1]
    new_section.left_margin=Inches(0)
    new_section.right_margin=Inches(0)
    new_section.top_margin=Inches(0)
    new_section.bottom_margin=Inches(0)
    doc.add_picture(image_path, width=new_section.page_width, height=10089400)
    doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # remaining page start
    new_section=doc.add_section()
    new_section.left_margin=Inches(0.3)
    new_section.right_margin=Inches(0.3)
    new_section.top_margin=Inches(0.3)
    new_section.bottom_margin=Inches(0.1)  

    new_section.start_type = WD_SECTION.CONTINUOUS
    header = new_section.header
    header.is_linked_to_previous = False
    header_paragraph = header.add_paragraph()
    # Add the header text to the paragraph
    run1 = header_paragraph.add_run("\t\t\t\t\tPartsOpsDashboard")
    run1.font.color.rgb = RGBColor(50,200,10)
    run1.font.size = Pt(25)
    run1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = doc.sections[-1]
    # Create a footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add the current date and time to the footer
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer_paragraph.text = f"Report generated on : {current_time}"

    return doc

#This function is a pipeline that uses "add_offset_colour","convert_word_to_pdf","remove_page_" respectively.Returns path of final pdf file 
def make_pdf(doc_path):
    doc_path=add_offset_colour(doc_path)
    current_time = f'{datetime.now():%Y-%m-%d %H-%M-%S}'
    pdf_filename = f'KPI_Report_{current_time}.pdf'
    pdf_path = os.path.join(os.path.dirname(doc_path), pdf_filename)
    convert_word_to_pdf(doc_path, pdf_path)
    new_pdf_path=remove_page_(pdf_path)
    if os.path.exists(pdf_path):
       os.remove(pdf_path)

    return new_pdf_path