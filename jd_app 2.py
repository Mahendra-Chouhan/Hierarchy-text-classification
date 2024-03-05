from docx import Document
from docx.shared import Inches
import streamlit as st
import pandas as pd
from PIL import Image
from io import BytesIO
import base64
import matplotlib.pyplot as plt
from lida import Manager, TextGenerationConfig, llm
from dotenv import load_dotenv
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import comtypes.client
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #added
from docx import Document
from datetime import datetime
from docx.shared import RGBColor  
import os
import openai
import tempfile
from docx.enum.section import WD_SECTION
from jd_utils import make_document_first_page,make_pdf


# Initialize the list to store file paths of visualizations
visualization_paths = []
st.set_page_config(page_title="Data Visualization", page_icon="📊")
st.write("# Auto Report📊")
# Load environment variables for API keys
load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')
# Initialize visualizations in session state
if 'visualization_paths' not in st.session_state:
   st.session_state['visualization_paths'] = []

if 'last_kpi_query' not in st.session_state:
    st.session_state['last_kpi_query'] = ""

if 'visualizations' not in st.session_state:
    st.session_state['visualizations'] = [] 

if "summary" not in st.session_state:
    st.session_state.summary = ""  # Initialize to an empty string 

if 'visual_description' not in st.session_state:
    st.session_state['visual_description'] = []     

datasets = [
   {"label": "CIDDS", "url": "https://raw.githubusercontent.com/shrinaneema98266/data/main/CIDDS_Preprocessed.csv"}]

# Function to convert base64 to Image object
def base64_to_image(base64_string):
   byte_data = base64.b64decode(base64_string)
   return Image.open(BytesIO(byte_data))

# Function to generate a textual summary using OpenAI
def generate_textual_summary(df):
   summary_data = df.head().to_csv(index=False)
   messages = [
       {"role": "system", "content": "Please summarize the following dataset."},
       {"role": "user", "content": summary_data}]
   model="gpt-3.5-turbo-0301"
   response = openai.chat.completions.create(model=model,messages=messages,temperature=0.5,max_tokens=500)
   return response.choices[0].message.content

def save_visualization_to_file(visualization_base64, index):
   image_data = base64.b64decode(visualization_base64)
   image_path = f'visualization{index}.png'
   with open(image_path, 'wb') as file:
       file.write(image_data)
   return image_path

# # Function to generate a natural language explanation of the visualization
def generate_visualization_explanation(chart_code):
   explanation = lida.explain(code=chart_code)
   return explanation
# Function to process and visualize KPIs with explanations
def process_visualization(kpi_query, index, display_image=True):
   visualization = lida.visualize(summary=st.session_state['summary'], goal=kpi_query, textgen_config=textgen_config, library='matplotlib')
   if visualization:
       image_base64 = visualization[0].raster
       img_path = save_visualization_to_file(image_base64, index)
       st.session_state['visualization_paths'].append(img_path)
       if display_image:
           img = Image.open(img_path)
           st.image(img)  # Display the image in the Streamlit app
           st.markdown("<br>", unsafe_allow_html = True)
           with st.expander("What the Visualization Tells Us"):
               visual_description = generate_visualization_explanation(visualization[0].code)
               st.session_state["visual_description"].append(visual_description[0])
               for description in visual_description[0]:
                   st.markdown(f"""**{description['section']}**""")
                   st.write(description["explanation"])
                   st.divider()
         
def add_selections_summary_kpis(doc, tenant_name, applications, start_date, end_date, summary, kpi_texts):
   
   p = doc.add_paragraph()
   run = p.add_run(f'Tenant Name : {tenant_name}')
   run.bold = True
   #run.font.color.rgb = RGBColor(50,200,10)
   p = doc.add_paragraph()
   run = p.add_run(f'Selected Applications : {", ".join(applications)}')
   run.bold = True
   #run.font.color.rgb = RGBColor(50,200,10)
   p = doc.add_paragraph()
   run = p.add_run(f'Start Date : {start_date}')
   run.bold = True
   #run.font.color.rgb = RGBColor(50,200,10)
   p = doc.add_paragraph()
   run = p.add_run(f'End Date : {end_date}')
   run.bold = True
   #run.font.color.rgb = RGBColor(50,200,10)
   # Space between end_date and summary
   doc.add_paragraph(' ')
   summary_para = doc.add_paragraph()
   summary_para.add_run("Summary:").bold = True
   for run in summary_para.runs:
       run.font.color.rgb = RGBColor(50,200,10)  # Set color to green
       run.font.size = Pt(15)
   doc.add_paragraph(textual_summary)
   doc.add_paragraph("")
   # Space between summary and KPIs
   kpi_heading = doc.add_paragraph()
   kpi_heading.add_run("KPI Report:").bold = True
   for run in kpi_heading.runs:
       run.font.color.rgb = RGBColor(50,200,10)  # Set color to green
       run.font.size = Pt(15)

# Initialize LIDA manager with text generation configuration
lida = Manager(text_gen=llm("openai"))
textgen_config = TextGenerationConfig(n=1, temperature=0.5, model="gpt-4", use_cache=True)
# Sidebar for application and tenant input
application_options = ['SAP Worksoft', 'Application 2', 'Application 3']
selected_applications = st.sidebar.multiselect('Select Applications', application_options)
# Sidebar for tenant input and file upload
tenant_name_selection = st.sidebar.selectbox('Select Tenant Name', datasets, format_func=lambda x: x['label'])
start_date = st.sidebar.date_input('📅 Start Date', key="start_date")
end_date = st.sidebar.date_input(' 📅 End Date', key="end_date")
if st.sidebar.button('Submit', key="submit"):
   st.session_state['visualization_paths'] = []
   st.session_state['submit_pressed'] = True
# Load data and generate summary when 'Submit' is pressed
if st.session_state.get('submit_pressed') and tenant_name_selection.get('url'):
   if tenant_name_selection['url'].endswith('.csv'):
       st.session_state['df'] = pd.read_csv(tenant_name_selection['url'])
   elif tenant_name_selection['url'].endswith('.json'):
       st.session_state['df'] = pd.read_json(tenant_name_selection['url'])
   st.dataframe(st.session_state['df'].head())
   st.subheader("Summary")
   try:
       textual_summary = generate_textual_summary(st.session_state['df'])
       st.write(textual_summary)
   except openai.OpenAIError as e:
       st.error(f"An error occurred while generating the summary: {str(e)}")
   st.session_state['summary'] = lida.summarize(st.session_state['df'], summary_method="default", textgen_config=textgen_config)
   st.session_state['processed_custom_kpi'] = False
   if 'processed_custom_kpis' not in st.session_state:
      st.subheader("KPIs")
     
      custom_kpis = ["What is the most common Class used in the network traffic?", "Generate a bar graph to show day trend of protocol?"]

      for i, kpi_query in enumerate(custom_kpis):
       #st.write(f"📊  {kpi_query}")
       st.markdown(f"#### 📊 {kpi_query}")
       process_visualization(kpi_query, i)
       st.session_state['processed_custom_kpis'] = True

def add_border_to_table_cell(cell, border_color="#808080", border_width="5"):
   # This function adds a border to a table cell
   sides = ['top', 'left', 'bottom', 'right']
   for side in sides:
       tag = 'w:{}'.format(side)
       border_elm = OxmlElement(tag)
       border_elm.set(qn('w:val'), 'single')
       border_elm.set(qn('w:sz'), border_width)
       border_elm.set(qn('w:color'), border_color)
       border_elm.set(qn('w:space'), "0")
       cell._tc.get_or_add_tcPr().append(border_elm)     

def add_visualizations_to_doc(document, visualizations):
    max_len=len(visualizations)-1
    cnt=0
    for img_path in visualizations:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            img = Image.open(img_path)
            img.save(tmp_file.name)        
            document.add_picture(tmp_file.name, width=Inches(3.2))
            if cnt<(max_len):
                doc.add_paragraph("________________________________________________________________________________________________________")
                
            cnt+=1
            
def add_visualizations_to_doc_1(document, visualizations,  max_columns=2):
   num_visualizations = len(visualizations)
   num_rows = (num_visualizations + max_columns - 1) // max_columns
   table = document.add_table(rows=num_rows, cols=max_columns)
   for i in range(num_rows):
       for j in range(max_columns):
           idx = i * max_columns + j
           if idx < num_visualizations:
               cell = table.cell(i, j)
               cell.margin_top = Inches(0.1)
               cell.margin_bottom = Inches(0.1)
               cell.margin_left = Inches(0.1)
               cell.margin_right = Inches(0.1)
               add_border_to_table_cell(cell, "#808080", "5")  # Set border color and width
               paragraph = cell.paragraphs[0]
               run = paragraph.add_run()
               run.add_picture(visualizations[idx], width=Inches(3.0))  # Set image size
           else:
               # If there are no more visualizations to add, simply continue
               continue   
       
# Function to edit visualization based on natural language instructions
def vis_edit(chart_code, summary, instructions, library, textgen_config):
   edited_charts = lida.edit(code=chart_code, summary=summary, instructions=instructions, library=library, textgen_config=textgen_config)
   return edited_charts[0]  # Return the first edited chart

def display_visualizations():
   for img_path in st.session_state['visualization_paths']:
      img = Image.open(img_path)
      st.image(img)
   
with st.sidebar.form(key="custom_kpis"):
    st.write("Custom KPI")
    input_lines = st.text_area(
        "Enter the KPI to generate visualization", height=50)
    KPI_submit_button = st.form_submit_button(label="Create")
    KPI_edit_button = st.sidebar.button("Edit Last Query")
    
if KPI_edit_button:
   # Load the last query into the text area for editing
   st.session_state["last_kpi_query"] = st.session_state["last_kpi_query"]
   st.experimental_rerun()      
    
# Logic for processing custom KPI and applying editing instructions
if KPI_submit_button and input_lines:
   user_input = input_lines.split("\n")
   kpi_query = user_input[0]  
   editing_instructions = user_input[1:]
   
   new_index = len(st.session_state['visualization_paths'])  # Index for the new visualization
   process_visualization(kpi_query, new_index, display_image=False)
   # Process editing instructions if provided
   if editing_instructions and 'summary' in st.session_state:
       # Assuming the first visualization's code is needed for editing
       initial_visualizations = lida.visualize(summary=st.session_state['summary'], goal=kpi_query, textgen_config=textgen_config, library='matplotlib')
       if initial_visualizations:
           initial_chart_code = initial_visualizations[0].code
           edited_chart = vis_edit(initial_chart_code, st.session_state['summary'], [editing_instructions], 'matplotlib', textgen_config)
           if edited_chart:
               image_base64 = edited_chart.raster
               img_path = save_visualization_to_file(image_base64, new_index)
               st.session_state['visualization_paths'].append(img_path)
               display_visualizations()
           else:
               st.error("No edited visualizations generated.")
       else:
           st.error("No initial visualizations generated for the KPI.")
   display_visualizations()
   initial_visualizations = lida.visualize(summary=st.session_state['summary'], goal=kpi_query, textgen_config=textgen_config, library='matplotlib')

   with st.expander("What the Visualization Tells Us"):
               visual_description = generate_visualization_explanation(initial_visualizations[0].code)
               st.session_state["visual_description"].append(visual_description[0])
               for description in visual_description[0]:
                   st.markdown(f"""**{description['section']}**""")
                   st.write(description["explanation"])
                   st.divider()   


# Generate the Word document when 'Generate Report' is clicked
if st.sidebar.button('Generate Report'):
   
   doc = Document()
   image_path=r"D:\JD_Support\JD_dashboard\Final_dash\Final_final_dash\report_dash\img8.jpg"
   Message = "Worksoft CTM"
   doc=make_document_first_page(doc,image_path,start_date,end_date,Message)

   #Retrieve user selections, summary, and KPIs
   tenant_name = tenant_name_selection.get('label', 'Unknown Tenant')  # Update as needed
   applications = selected_applications
   start_date = st.session_state.get("start_date", "Not Specified")
   end_date = st.session_state.get("end_date", "Not Specified")
   summary = st.session_state.get('summary', "No summary available.")

   kpi_texts = [kpi["text"] for kpi in st.session_state.get('kpis', [])]  # Assuming KPIs are stored in session state
   add_selections_summary_kpis(doc, tenant_name, applications, start_date, end_date, summary, kpi_texts)
   add_visualizations_to_doc(doc,st.session_state['visualization_paths'])
   
   doc_path = f'KPI_Report_{datetime.now():%Y-%m-%d_%H-%M-%S}.docx'
   doc.save(doc_path)
   new_pdf_path=make_pdf(doc_path)
   display_visualizations()

   # Open the PDF file and create a download button
   with open(new_pdf_path, 'rb') as pdf_file:
           pdf_data = pdf_file.read()
           b64_pdf = base64.b64encode(pdf_data).decode('utf-8')
           href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="KPI_Report_s.pdf">Download PDF Document</a>'
           st.markdown(href, unsafe_allow_html=True)
           
   for path in visualization_paths:
           os.remove(path)
   os.remove(doc_path)

   
  




   


